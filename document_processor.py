"""
Document processing module for loading, chunking, and preparing documents.

Supports multiple document formats with extensible loader architecture.
"""

import hashlib
import inspect
import logging
import os
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Union
import mimetypes

import pandas as pd
from langchain_community.document_loaders import (
    TextLoader,
    PyPDFLoader,
    WebBaseLoader,
    UnstructuredFileLoader
)
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from config import get_config

# OCR support (optional - graceful fallback if not installed)
try:
    import pytesseract
    from PIL import Image
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    pytesseract = None
    Image = None
    convert_from_path = None

try:
    from desktop_app.utils.hashing import calculate_file_hash
except ImportError:
    # Fallback if desktop_app not available (e.g. strict server environment)
    import xxhash
    def calculate_file_hash(path: Path) -> str:
        hasher = xxhash.xxh64()
        with open(path, 'rb') as f:
            while chunk := f.read(8192):
                hasher.update(chunk)
        return hasher.hexdigest()

logger = logging.getLogger(__name__)


class DocumentProcessingError(Exception):
    """Base exception for document processing errors."""
    pass


class UnsupportedFormatError(DocumentProcessingError):
    """Exception for unsupported document formats."""
    pass


class LoaderError(DocumentProcessingError):
    """Exception for document loading errors."""
    pass


@dataclass
class ProcessedDocument:
    """Container for processed document data."""
    
    document_id: str
    source_uri: str
    chunks: List[Document]
    metadata: Dict[str, Any]
    processed_at: datetime
    
    def __len__(self) -> int:
        """Return number of chunks."""
        return len(self.chunks)
    
    def get_chunk_texts(self) -> List[str]:
        """Get list of chunk texts."""
        return [chunk.page_content for chunk in self.chunks]


class DocumentLoader:
    """Base class for document loaders."""
    
    def can_load(self, source_uri: str) -> bool:
        """Check if this loader can handle the source."""
        raise NotImplementedError
    
    def load(self, source_uri: str) -> List[Document]:
        """Load document and return list of Document objects."""
        raise NotImplementedError
    
    def get_metadata(self, source_uri: str) -> Dict[str, Any]:
        """Extract metadata from source."""
        return {}


class TextDocumentLoader(DocumentLoader):
    """Loader for plain text files and Markdown."""
    
    def can_load(self, source_uri: str) -> bool:
        """Check if source is a text, Markdown, YAML, or allowed config file."""
        path = Path(source_uri)
        is_text_ext = path.suffix.lower() in ['.txt', '.md', '.markdown', '.yaml', '.yml']
        # We can't easily access config here without passing it down, 
        # but we can check common text filenames
        is_text_filename = path.name in ['LICENSE', 'Dockerfile', 'Makefile', 'Jenkinsfile']
        return is_text_ext or is_text_filename
    
    def load(self, source_uri: str) -> List[Document]:
        """Load text file with encoding detection."""
        try:
            # Try default utf-8 first
            loader = TextLoader(source_uri, encoding='utf-8')
            return loader.load()
        except Exception:
            try:
                # Try autodetect encoding
                loader = TextLoader(source_uri, autodetect_encoding=True)
                return loader.load()
            except Exception:
                try:
                    # Fallback to latin-1 (common for Windows logs)
                    loader = TextLoader(source_uri, encoding='latin-1')
                    return loader.load()
                except Exception as e:
                    logger.error(f"Failed to load text file {source_uri}: {e}")
                    raise LoaderError(f"Text loading failed: {e}")
    
    def get_metadata(self, source_uri: str) -> Dict[str, Any]:
        """Get text file metadata."""
        path = Path(source_uri)
        return {
            'file_type': 'text',
            'file_size': path.stat().st_size if path.exists() else 0,
            'file_extension': path.suffix
        }


class PDFDocumentLoader(DocumentLoader):
    """Loader for PDF files with OCR fallback for scanned documents."""
    
    # Image extensions we support for OCR
    IMAGE_EXTENSIONS = ['.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp']
    
    def can_load(self, source_uri: str) -> bool:
        """Check if source is a PDF file."""
        return source_uri.lower().endswith('.pdf')
    
    def load(self, source_uri: str, ocr_mode: str = 'auto') -> List[Document]:
        """
        Load PDF file with optional OCR fallback.
        
        Args:
            source_uri: Path to PDF file
            ocr_mode: 'skip' (no OCR), 'only' (OCR only), 'auto' (native first, OCR fallback)
        """
        try:
            # Try native text extraction first (unless ocr_mode is 'only')
            if ocr_mode != 'only':
                loader = PyPDFLoader(source_uri)
                documents = loader.load()
                
                # Check if we got meaningful text
                total_text = ''.join(doc.page_content for doc in documents)
                if total_text.strip() and len(total_text.strip()) > 50:
                    logger.debug(f"PDF has native text ({len(total_text)} chars)")
                    return documents
                
                # If skip mode, return what we have (even if empty)
                if ocr_mode == 'skip':
                    logger.info(f"PDF has little/no text, OCR skipped (mode=skip)")
                    return documents
            
            # OCR fallback
            if not OCR_AVAILABLE:
                if ocr_mode == 'only':
                    raise LoaderError("OCR mode=only but pytesseract not installed")
                logger.warning(f"PDF appears scanned but OCR not available")
                return documents if ocr_mode != 'only' else []
            
            logger.info(f"Attempting OCR extraction for PDF: {source_uri}")
            return self._extract_with_ocr(source_uri)
            
        except Exception as e:
            logger.error(f"Failed to load PDF {source_uri}: {e}")
            raise LoaderError(f"PDF loading failed: {e}")
    
    def _extract_with_ocr(self, source_uri: str) -> List[Document]:
        """Extract text from PDF using OCR."""
        config = get_config()
        documents = []
        
        try:
            # Convert PDF pages to images
            images = convert_from_path(
                source_uri,
                dpi=config.ocr.dpi,
                fmt='png'
            )
            
            for page_num, image in enumerate(images, start=1):
                # Run OCR on each page
                text = pytesseract.image_to_string(
                    image,
                    lang=config.ocr.language,
                    timeout=config.ocr.timeout
                )
                
                if text.strip():
                    documents.append(Document(
                        page_content=text,
                        metadata={
                            'source': source_uri,
                            'page': page_num,
                            'extraction_method': 'ocr'
                        }
                    ))
            
            logger.info(f"OCR extracted {len(documents)} pages from PDF")
            return documents
            
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            raise LoaderError(f"OCR extraction failed: {e}")
    
    def get_metadata(self, source_uri: str) -> Dict[str, Any]:
        """Get PDF metadata."""
        path = Path(source_uri)
        return {
            'file_type': 'pdf',
            'file_size': path.stat().st_size if path.exists() else 0,
            'file_extension': '.pdf'
        }


class ImageDocumentLoader(DocumentLoader):
    """Loader for image files using OCR."""
    
    SUPPORTED_EXTENSIONS = ['.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp']
    
    def can_load(self, source_uri: str) -> bool:
        """Check if source is an image file."""
        return any(source_uri.lower().endswith(ext) for ext in self.SUPPORTED_EXTENSIONS)
    
    def load(self, source_uri: str, ocr_mode: str = 'auto') -> List[Document]:
        """
        Load image file using OCR.
        
        Args:
            source_uri: Path to image file
            ocr_mode: 'skip' returns empty, 'only'/'auto' runs OCR
        """
        if ocr_mode == 'skip':
            logger.info(f"Skipping image (OCR mode=skip): {source_uri}")
            return []
        
        if not OCR_AVAILABLE:
            raise LoaderError(
                f"Cannot process image {source_uri}: pytesseract not installed. "
                "Install OCR support: pip install pytesseract Pillow"
            )
        
        try:
            config = get_config()
            image = Image.open(source_uri)
            
            # Run OCR
            text = pytesseract.image_to_string(
                image,
                lang=config.ocr.language,
                timeout=config.ocr.timeout
            )
            
            if not text.strip():
                logger.warning(f"No text extracted from image: {source_uri}")
                return []
            
            return [Document(
                page_content=text,
                metadata={
                    'source': source_uri,
                    'extraction_method': 'ocr',
                    'image_format': image.format,
                    'image_size': f"{image.width}x{image.height}"
                }
            )]
            
        except Exception as e:
            logger.error(f"Failed to OCR image {source_uri}: {e}")
            raise LoaderError(f"Image OCR failed: {e}")
    
    def get_metadata(self, source_uri: str) -> Dict[str, Any]:
        """Get image metadata."""
        path = Path(source_uri)
        ext = path.suffix.lower()
        return {
            'file_type': 'image',
            'file_size': path.stat().st_size if path.exists() else 0,
            'file_extension': ext
        }


class WebDocumentLoader(DocumentLoader):
    """Loader for web URLs."""
    
    def can_load(self, source_uri: str) -> bool:
        """Check if source is a web URL."""
        return source_uri.startswith(('http://', 'https://'))
    
    def load(self, source_uri: str) -> List[Document]:
        """Load web page."""
        try:
            loader = WebBaseLoader(source_uri)
            return loader.load()
        except Exception as e:
            logger.error(f"Failed to load web page {source_uri}: {e}")
            raise LoaderError(f"Web loading failed: {e}")
    
    def get_metadata(self, source_uri: str) -> Dict[str, Any]:
        """Get web page metadata."""
        return {
            'file_type': 'web',
            'url': source_uri
        }


class OfficeDocumentLoader(DocumentLoader):
    """Loader for Microsoft Office documents."""

    SUPPORTED_EXTENSIONS = ['.doc', '.docx', '.pptx', '.html']
    CONVERTER_CANDIDATES = ("soffice", "libreoffice")
    WINDOWS_DEFAULT_PATHS = (
        "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
        "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
    )

    def can_load(self, source_uri: str) -> bool:
        """Check if source is an Office document."""
        return any(source_uri.lower().endswith(ext) for ext in self.SUPPORTED_EXTENSIONS)
    
    def load(self, source_uri: str) -> List[Document]:
        """Load Office document based on extension."""
        lowered = source_uri.lower()
        if lowered.endswith(('.doc', '.docx')):
            return self._load_word_document(source_uri)
        if lowered.endswith('.html'):
            return self._load_doc_with_unstructured(source_uri)
        if lowered.endswith('.pptx'):
            loader = UnstructuredFileLoader(source_uri)
            return loader.load()
        raise UnsupportedFormatError(f"Unsupported office extension for loader: {source_uri}")
    
    def _load_doc_with_unstructured(self, source_uri: str) -> List[Document]:
        try:
            loader = UnstructuredFileLoader(source_uri)
            documents = loader.load()
            if documents:
                return documents
            raise LoaderError(
                "Legacy .doc format is not supported. Please convert the document to .docx before uploading."
            )
        except Exception as exc:
            logger.error(f"Fallback unstructured load failed for {source_uri}: {exc}")
            raise LoaderError(
                "Legacy .doc format is not supported. Please convert the document to .docx before uploading."
            ) from exc

    def _find_converter_command(self) -> Optional[str]:
        override = os.getenv("LIBREOFFICE_PATH")
        if override:
            if Path(override).is_file() or shutil.which(override):
                return override
        for candidate in self.CONVERTER_CANDIDATES:
            located = shutil.which(candidate)
            if located:
                return located

        for win_path in self.WINDOWS_DEFAULT_PATHS:
            converted = convert_windows_path(win_path)
            if Path(converted).exists():
                return converted
        return None

    def _convert_doc_to_docx(self, source_uri: str) -> Optional[Path]:
        command = self._find_converter_command()
        if not command:
            logger.warning("No LibreOffice/soffice command found for automatic .doc conversion")
            return None

        with tempfile.TemporaryDirectory(prefix="doc_convert_") as tmpdir:
            output_dir = Path(tmpdir)
            cmd = [
                command,
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(output_dir),
                source_uri
            ]
            try:
                subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            except Exception as exc:
                logger.error(f"Automatic .doc conversion via {command} failed for {source_uri}: {exc}")
                return None

            expected = output_dir / (Path(source_uri).stem + ".docx")
            if not expected.exists():
                logger.error(f"Conversion reported success but {expected} was not created")
                return None

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
                shutil.copyfile(expected, tmp_file.name)
                temp_path = tmp_file.name

        return Path(temp_path)

    def _extract_docx_documents(self, source_uri: str, *, original_source: Optional[str] = None) -> List[Document]:
        try:
            doc = DocxDocument(source_uri)
        except Exception as exc:
            logger.error(f"Failed to load Word document {source_uri}: {exc}")
            raise LoaderError(f"Word document loading failed: {exc}")

        text_segments: List[str] = []

        for paragraph in doc.paragraphs:
            cleaned = paragraph.text.strip()
            if cleaned:
                text_segments.append(cleaned)

        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_segments.append("\t".join(row_text))

        full_text = "\n\n".join(text_segments).strip()
        if not full_text:
            raise LoaderError("No textual content could be extracted from the Word document.")

        metadata_source = original_source or source_uri
        return [Document(page_content=full_text, metadata={'source': metadata_source})]

    def _load_word_document(self, source_uri: str, *, original_source: Optional[str] = None) -> List[Document]:
        try:
            return self._extract_docx_documents(source_uri, original_source=original_source)
        except LoaderError as exc:
            # Re-raise immediately for non-legacy doc
            if not source_uri.lower().endswith('.doc'):
                raise exc

            logger.warning("python-docx cannot open legacy .doc; attempting automatic conversion")
            converted_path = self._convert_doc_to_docx(source_uri)
            if converted_path:
                try:
                    return self._extract_docx_documents(str(converted_path), original_source=original_source or source_uri)
                finally:
                    try:
                        os.remove(converted_path)
                    except OSError:
                        pass

            # Fall back to unstructured if conversion failed
            return self._load_doc_with_unstructured(source_uri)

        try:
            loader = UnstructuredFileLoader(source_uri)
            return loader.load()
        except Exception as e:
            logger.error(f"Failed to load Office document {source_uri}: {e}")
            raise LoaderError(f"Office document loading failed: {e}")
    
    def get_metadata(self, source_uri: str) -> Dict[str, Any]:
        """Get Office document metadata."""
        path = Path(source_uri)
        return {
            'file_type': 'office',
            'file_size': path.stat().st_size if path.exists() else 0,
            'file_extension': path.suffix
        }


class SpreadsheetLoader(DocumentLoader):
    """Loader for spreadsheet files (Excel, CSV)."""
    
    SUPPORTED_EXTENSIONS = ['.xlsx', '.xls', '.csv']
    
    def can_load(self, source_uri: str) -> bool:
        """Check if source is a spreadsheet."""
        return any(source_uri.lower().endswith(ext) for ext in self.SUPPORTED_EXTENSIONS)
    
    def load(self, source_uri: str) -> List[Document]:
        """Load spreadsheet and convert to documents."""
        try:
            # Read spreadsheet
            if source_uri.lower().endswith('.csv'):
                df = pd.read_csv(source_uri)
            else:
                df = pd.read_excel(source_uri)
            
            # Convert rows to text documents
            documents = []
            for index, row in df.iterrows():
                # Create descriptive text from row
                text = f"Row {index + 1}: " + ", ".join(
                    [f"{col}: {val}" for col, val in row.items() if pd.notna(val)]
                )
                doc = Document(
                    page_content=text,
                    metadata={'row_index': index, 'source': source_uri}
                )
                documents.append(doc)
            
            return documents
        except Exception as e:
            logger.error(f"Failed to load spreadsheet {source_uri}: {e}")
            raise LoaderError(f"Spreadsheet loading failed: {e}")
    
    def get_metadata(self, source_uri: str) -> Dict[str, Any]:
        """Get spreadsheet metadata."""
        path = Path(source_uri)
        return {
            'file_type': 'spreadsheet',
            'file_size': path.stat().st_size if path.exists() else 0,
            'file_extension': path.suffix
        }


class DocumentProcessor:
    """
    Main document processor for loading, chunking, and preparing documents.
    
    Orchestrates document loading, text splitting, and metadata extraction.
    """
    
    def __init__(self):
        """Initialize document processor."""
        self.config = get_config()
        self.loaders: List[DocumentLoader] = [
            TextDocumentLoader(),
            PDFDocumentLoader(),
            ImageDocumentLoader(),  # OCR for images
            WebDocumentLoader(),
            OfficeDocumentLoader(),
            SpreadsheetLoader()
        ]
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.config.chunking.size,
            chunk_overlap=self.config.chunking.overlap,
            separators=self.config.chunking.separators,
            length_function=len
        )
    
    def _get_loader(self, source_uri: str) -> Optional[DocumentLoader]:
        """Get appropriate loader for source."""
        for loader in self.loaders:
            if loader.can_load(source_uri):
                return loader
        return None
    
    def _generate_document_id(self, source_uri: str) -> str:
        """Generate unique document ID from source URI."""
        # Use SHA256 hash of source URI for consistent IDs
        return hashlib.sha256(source_uri.encode()).hexdigest()[:16]
    
    def _validate_source(self, source_uri: str) -> None:
        """Validate source URI."""
        # Check if it's a web URL
        if source_uri.startswith(('http://', 'https://')):
            return
        
        # Check if it's a local file
        path = Path(source_uri)
        if not os.path.exists(source_uri):
            raise FileNotFoundError(f"Source file not found: {source_uri}")
        
        # Check file size
        if path.is_file():
            max_size = self.config.max_file_size_mb * 1024 * 1024
            if path.stat().st_size > max_size:
                raise DocumentProcessingError(
                    f"File too large: {path.stat().st_size} bytes "
                    f"(max: {max_size} bytes)"
                )
        
        # Check extension or filename
        is_supported_extension = path.suffix.lower() in self.config.supported_extensions
        is_supported_filename = path.name in self.config.supported_filenames
        
        if not (is_supported_extension or is_supported_filename):
            raise UnsupportedFormatError(
                f"Unsupported file: {path.name} (extension: {path.suffix})"
            )
    
    def process(
        self,
        source_uri: str,
        custom_metadata: Optional[Dict[str, Any]] = None,
        ocr_mode: Optional[str] = None
    ) -> ProcessedDocument:
        """
        Process document from source URI.
        
        Args:
            source_uri: Path or URL to document
            custom_metadata: Optional custom metadata to add
            ocr_mode: OCR processing mode ('skip', 'only', 'auto', or None for config default)
            
        Returns:
            ProcessedDocument with chunks and metadata
            
        Raises:
            DocumentProcessingError: If processing fails
            UnsupportedFormatError: If format is not supported
        """
        logger.info(f"Processing document: {source_uri}")
        
        # Use config default if not specified
        if ocr_mode is None:
            ocr_mode = self.config.ocr.mode
        
        # Validate source
        self._validate_source(source_uri)
        
        # Calculate file hash if it's a local file
        file_hash = None
        if not source_uri.startswith(('http://', 'https://')):
            file_hash = calculate_file_hash(Path(source_uri))
        
        # Get appropriate loader
        loader = self._get_loader(source_uri)
        if not loader:
            raise UnsupportedFormatError(
                f"No loader available for: {source_uri}"
            )
        
        # Load document (pass ocr_mode if loader supports it)
        try:
            # Check if loader accepts ocr_mode parameter
            sig = inspect.signature(loader.load)
            if 'ocr_mode' in sig.parameters:
                documents = loader.load(source_uri, ocr_mode=ocr_mode)
            else:
                documents = loader.load(source_uri)
            
            # Sanitize content (remove null bytes which Postgres rejects)
            for doc in documents:
                if doc.page_content:
                    doc.page_content = doc.page_content.replace('\x00', '')

            if not documents:
                raise LoaderError("No content loaded from document")
        except DocumentProcessingError:
            raise
        except LoaderError as e:
            # Preserve loader-specific message (e.g., legacy .doc conversion hint)
            raise e
        except Exception as e:
            logger.error(f"Failed to load document: {e}")
            cause = getattr(e, "__cause__", None) or getattr(e, "__context__", None)
            if isinstance(cause, LoaderError):
                raise cause
            raise DocumentProcessingError(f"Document loading failed: {e}") from e
        
        # Split into chunks
        try:
            chunks = self.text_splitter.split_documents(documents)
            logger.info(f"Split document into {len(chunks)} chunks")
        except Exception as e:
            logger.error(f"Failed to split document: {e}")
            raise DocumentProcessingError(f"Chunking failed: {e}")
        
        # Generate document ID
        document_id = self._generate_document_id(source_uri)
        
        # Collect metadata
        metadata = loader.get_metadata(source_uri)
        metadata.update({
            'document_id': document_id,
            'source_uri': source_uri,
            'chunk_count': len(chunks),
            'chunking_config': {
                'chunk_size': self.config.chunking.size,
                'chunk_overlap': self.config.chunking.overlap
            },
            'file_hash': file_hash,
            'processed_at': datetime.utcnow().isoformat()
        })
        
        # Add custom metadata
        if custom_metadata:
            metadata.update(custom_metadata)
        
        return ProcessedDocument(
            document_id=document_id,
            source_uri=source_uri,
            chunks=chunks,
            metadata=metadata,
            processed_at=datetime.utcnow()
        )
    
    def process_batch(
        self,
        source_uris: List[str],
        custom_metadata: Optional[Dict[str, Any]] = None
    ) -> List[ProcessedDocument]:
        """
        Process multiple documents.
        
        Args:
            source_uris: List of paths or URLs
            custom_metadata: Optional custom metadata for all documents
            
        Returns:
            List of ProcessedDocument objects
        """
        processed_docs = []
        
        for source_uri in source_uris:
            try:
                doc = self.process(source_uri, custom_metadata)
                processed_docs.append(doc)
            except Exception as e:
                logger.error(f"Failed to process {source_uri}: {e}")
                # Continue with other documents
        
        return processed_docs


def convert_windows_path(path: str) -> str:
    """
    Convert Windows path to WSL path if needed.
    
    Args:
        path: Windows or WSL path
        
    Returns:
        WSL-compatible path
    """
    # Handle UNC paths pointing to WSL
    if path.startswith("\\\\wsl.localhost\\"):
        no_prefix = path.replace("\\wsl.localhost\\", "", 1)
        unix_path = no_prefix.replace('\\', '/')
        logger.info(f"Converted WSL UNC path: {path} -> /{unix_path.lstrip('/')}")
        return "/" + unix_path.lstrip('/')

    # Check if it's a Windows path (e.g., C:\...)
    if len(path) > 1 and path[1] == ':' and path[0].isalpha():
        drive_letter = path[0].lower()
        # Convert C:\Users\... to /mnt/c/Users/...
        wsl_path = f"/mnt/{drive_letter}/{path[3:].replace(chr(92), '/')}"
        logger.info(f"Converted Windows path: {path} -> {wsl_path}")
        return wsl_path

    return path
